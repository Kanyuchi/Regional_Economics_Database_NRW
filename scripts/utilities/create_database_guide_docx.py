"""
Create Professional Database Structure Guide (Word Document)
For Regional Economics Database for NRW
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    return heading

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def main():
    """Create the Word document."""
    
    doc = Document()
    
    # ========================================
    # TITLE PAGE
    # ========================================
    title = doc.add_heading('Regional Economics Database for NRW', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Database Structure and Design Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Document info
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light List Accent 1'
    
    info_data = [
        ('Version:', '1.0'),
        ('Database Name:', 'regional_economics'),
        ('Architecture:', 'Star Schema (Data Warehouse)'),
        ('Created:', 'December 2024'),
        ('Last Updated:', 'December 17, 2025')
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        info_table.rows[i].cells[1].text = value
    
    add_page_break(doc)
    
    # ========================================
    # EXECUTIVE SUMMARY
    # ========================================
    add_styled_heading(doc, 'Executive Summary', 1)
    
    doc.add_paragraph(
        'This database stores economic, demographic, and labor market indicators for '
        'North Rhine-Westphalia (NRW) regions using a star schema design. All metrics '
        'from 36+ different source tables are stored in a unified structure with shared '
        'dimension tables for geography, time, and indicators.'
    )
    
    doc.add_paragraph('Key Benefits:', style='Heading 3')
    benefits = doc.add_paragraph()
    benefits.style = 'List Bullet'
    benefits.text = 'Single query pattern works for all data types'
    
    for benefit in [
        'Easy to add new indicators without schema changes',
        'Optimized for analytical queries and reporting',
        'Consistent data structure across all categories',
        'Scales efficiently with growing data volumes'
    ]:
        p = doc.add_paragraph(benefit, style='List Bullet')
    
    # ========================================
    # DATABASE ARCHITECTURE
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '1. Database Architecture', 1)
    
    add_styled_heading(doc, 'Star Schema Overview', 2)
    
    doc.add_paragraph(
        'The database follows a star schema pattern where dimension tables (reference data) '
        'surround fact tables (measurements). This design optimizes analytical queries and '
        'provides a consistent structure across different data types.'
    )
    
    # Architecture diagram as text
    doc.add_paragraph('Architecture Diagram:', style='Heading 3')
    
    arch_para = doc.add_paragraph()
    arch_para.style = 'No Spacing'
    arch_para.paragraph_format.space_after = Pt(0)
    
    # Add preformatted text
    arch_text = """
    DIMENSION TABLES (Reference/Lookup Data)
    ═══════════════════════════════════════════
    
    ┌─────────────────┐         ┌─────────────────┐
    │ dim_geography   │         │    dim_time     │
    ├─────────────────┤         ├─────────────────┤
    │ • geo_id (PK)   │         │ • time_id (PK)  │
    │ • region_code   │         │ • year          │
    │ • region_name   │         │ • reference_date│
    │ • region_type   │         │ • quarter       │
    │ • ruhr_area     │         └─────────────────┘
    └─────────────────┘
    
          ┌─────────────────┐
          │  dim_indicator  │
          ├─────────────────┤
          │ • indicator_id  │
          │ • indicator_code│
          │ • category      │
          │ • source_table  │
          └─────────────────┘
    
    
    FACT TABLES (Measurement Data)
    ═══════════════════════════════
    
    ┌──────────────────────────────────────┐
    │      fact_demographics               │
    │  (Population & Employment Data)      │
    ├──────────────────────────────────────┤
    │ • geo_id (FK) → dim_geography       │
    │ • time_id (FK) → dim_time           │
    │ • indicator_id (FK) → dim_indicator │
    │ • value (THE MEASUREMENT)           │
    │ • gender, nationality, age_group    │
    └──────────────────────────────────────┘
    """
    
    for line in arch_text.split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        if p.runs:
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
    
    # ========================================
    # DIMENSION TABLES
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '2. Dimension Tables', 1)
    
    doc.add_paragraph(
        'Dimension tables contain descriptive attributes and reference data that provide '
        'context for the measurements in fact tables.'
    )
    
    # dim_geography
    add_styled_heading(doc, '2.1 dim_geography (Geographic Dimension)', 2)
    
    doc.add_paragraph('Purpose: Defines all geographic regions in the database')
    doc.add_paragraph('Current Count: 60 regions')
    
    geo_headers = ['Column', 'Type', 'Description', 'Example']
    geo_rows = [
        ('geo_id', 'SERIAL (PK)', 'Unique identifier', '1, 2, 3...'),
        ('region_code', 'VARCHAR(20)', 'Official region code', '05112, 05, DG'),
        ('region_name', 'VARCHAR(255)', 'Region name (German)', 'Duisburg, NRW'),
        ('region_type', 'VARCHAR(50)', 'Type of region', 'urban_district, state'),
        ('ruhr_area', 'BOOLEAN', 'Part of Ruhr?', 'TRUE/FALSE'),
        ('latitude', 'DECIMAL', 'GPS coordinate', '51.4344'),
        ('longitude', 'DECIMAL', 'GPS coordinate', '6.7623'),
        ('area_sqkm', 'DECIMAL', 'Area in km²', '232.82'),
    ]
    
    create_table(doc, geo_headers, geo_rows)
    
    doc.add_paragraph()
    doc.add_paragraph('Region Types:', style='Heading 3')
    types_list = [
        'district - Rural district (Kreis)',
        'urban_district - Independent city (Kreisfreie Stadt)',
        'administrative_district - Regional government area (Regierungsbezirk)',
        'state - Federal state (Bundesland)',
        'country - National level (Deutschland)'
    ]
    for item in types_list:
        doc.add_paragraph(item, style='List Bullet')
    
    # dim_time
    add_styled_heading(doc, '2.2 dim_time (Temporal Dimension)', 2)
    
    doc.add_paragraph('Purpose: Defines temporal periods for data')
    doc.add_paragraph('Current Count: 17 years (2008-2024)')
    
    time_headers = ['Column', 'Type', 'Description', 'Example']
    time_rows = [
        ('time_id', 'SERIAL (PK)', 'Unique identifier', '1, 2, 3...'),
        ('year', 'INTEGER', 'Calendar year', '2024'),
        ('reference_date', 'DATE', 'Specific date', '2024-06-30'),
        ('reference_type', 'VARCHAR(50)', 'Type of reference', 'mid_year'),
        ('quarter', 'INTEGER', 'Quarter (1-4)', '2'),
        ('month', 'INTEGER', 'Month (1-12)', '6'),
    ]
    
    create_table(doc, time_headers, time_rows)
    
    # dim_indicator
    add_page_break(doc)
    add_styled_heading(doc, '2.3 dim_indicator (Indicator Dimension)', 2)
    
    doc.add_paragraph('Purpose: Defines what each measurement represents')
    doc.add_paragraph('Current Count: 4 indicators (40+ planned)')
    
    ind_headers = ['Column', 'Type', 'Description']
    ind_rows = [
        ('indicator_id', 'SERIAL (PK)', 'Unique identifier'),
        ('indicator_code', 'VARCHAR(100)', 'Short code (e.g., pop_total)'),
        ('indicator_name', 'VARCHAR(255)', 'Full name (German)'),
        ('indicator_category', 'VARCHAR(100)', 'Category (e.g., demographics)'),
        ('source_table_id', 'VARCHAR(50)', 'GENESIS table (e.g., 12411-03-03-4)'),
        ('unit_of_measure', 'VARCHAR(50)', 'Unit (e.g., persons, employees)'),
        ('update_frequency', 'VARCHAR(50)', 'How often updated (e.g., annual)'),
    ]
    
    create_table(doc, ind_headers, ind_rows)
    
    doc.add_paragraph()
    doc.add_paragraph('Currently Loaded Indicators:', style='Heading 3')
    
    loaded_headers = ['ID', 'Code', 'Name', 'Source Table', 'Records']
    loaded_rows = [
        ('1', 'pop_total', 'Population total', '12411-03-03-4', '17,556'),
        ('2', 'employment_workplace', 'Employment at workplace', '13111-01-03-4', '798'),
        ('9', 'employment_sector', 'Employment by sector', '13111-07-05-4', '19,134'),
        ('3', 'employment_scope', 'Employment by scope', '13111-03-02-4', '~8,700'),
    ]
    
    create_table(doc, loaded_headers, loaded_rows)
    
    # ========================================
    # FACT TABLES
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '3. Fact Tables', 1)
    
    doc.add_paragraph(
        'Fact tables contain measurements (the actual numbers) with foreign keys linking '
        'to dimension tables. Each record represents a specific measurement for a particular '
        'region, time period, and indicator.'
    )
    
    add_styled_heading(doc, '3.1 fact_demographics', 2)
    
    doc.add_paragraph('Purpose: Population and demographic indicators')
    doc.add_paragraph('Current Records: 45,000+')
    
    demo_headers = ['Column', 'Type', 'Description']
    demo_rows = [
        ('geo_id', 'INTEGER (FK)', 'Links to dim_geography → Which region?'),
        ('time_id', 'INTEGER (FK)', 'Links to dim_time → Which year/period?'),
        ('indicator_id', 'INTEGER (FK)', 'Links to dim_indicator → What metric?'),
        ('value', 'NUMERIC(20,4)', 'THE MEASUREMENT - the actual number'),
        ('gender', 'VARCHAR(20)', 'male, female, total'),
        ('nationality', 'VARCHAR(50)', 'german, foreign, total'),
        ('age_group', 'VARCHAR(50)', '0-5, 6-17, 18-64, 65+, total'),
        ('notes', 'TEXT', 'Additional info (sector, scope)'),
        ('data_quality_flag', 'VARCHAR(20)', 'V=Validated, E=Estimated, P=Provisional'),
    ]
    
    create_table(doc, demo_headers, demo_rows)
    
    doc.add_paragraph()
    doc.add_paragraph('Example Record:', style='Heading 3')
    
    example_text = """
    geo_id = 5           → Duisburg (from dim_geography)
    time_id = 15         → Year 2024, June 30 (from dim_time)
    indicator_id = 9     → Employment by sector (from dim_indicator)
    value = 156,999.00   → THE ACTUAL NUMBER
    gender = 'total'
    nationality = 'total'
    notes = 'Sector: Dienstleistungsbereiche (G-U)'
    
    INTERPRETATION: In Duisburg on June 30, 2024, there were 
    156,999 employees in the service sector.
    """
    
    example_para = doc.add_paragraph()
    example_para.style = 'Quote'
    run = example_para.add_run(example_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    
    # ========================================
    # HOW IT WORKS
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '4. How It Works', 1)
    
    add_styled_heading(doc, 'Traditional vs Star Schema Approach', 2)
    
    doc.add_paragraph('Traditional Approach (Complex):')
    trad_list = [
        'tbl_population → Unique structure',
        'tbl_employment → Different structure',
        'tbl_unemployment → Different structure',
        'tbl_gdp → Different structure',
        '... 36 different table structures with different query patterns'
    ]
    for item in trad_list:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph('Star Schema Approach (Unified):')
    star_list = [
        'ALL data → fact_demographics, fact_labor_market, etc.',
        'Same query pattern for everything',
        'indicator_id tells you what the data means',
        'Add new indicators without changing schema'
    ]
    for item in star_list:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(0, 150, 0)
    
    # ========================================
    # EXAMPLE QUERIES
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '5. Example Queries', 1)
    
    add_styled_heading(doc, '5.1 Get Duisburg Population for 2024', 2)
    
    query1 = """
SELECT 
    g.region_name,
    t.year,
    i.indicator_name,
    f.value
FROM fact_demographics f
JOIN dim_geography g ON f.geo_id = g.geo_id
JOIN dim_time t ON f.time_id = t.time_id
JOIN dim_indicator i ON f.indicator_id = i.indicator_id
WHERE g.region_code = '05112'           -- Duisburg
  AND t.year = 2024
  AND i.indicator_code = 'pop_total';
    """
    
    query_para = doc.add_paragraph()
    query_para.style = 'Quote'
    run = query_para.add_run(query1)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 200)
    
    doc.add_paragraph('Result:')
    result_headers = ['region_name', 'year', 'indicator_name', 'value']
    result_rows = [('Duisburg', '2024', 'Bevölkerung insgesamt', '502,270')]
    create_table(doc, result_headers, result_rows)
    
    # Query 2
    add_styled_heading(doc, '5.2 Employment Trend for Duisburg (2020-2024)', 2)
    
    query2 = """
SELECT 
    t.year,
    SUM(f.value) as total_employment
FROM fact_demographics f
JOIN dim_geography g ON f.geo_id = g.geo_id
JOIN dim_time t ON f.time_id = t.time_id
WHERE g.region_code = '05112'           -- Duisburg
  AND f.indicator_id = 9                 -- Employment by sector
  AND t.year BETWEEN 2020 AND 2024
GROUP BY t.year
ORDER BY t.year;
    """
    
    query_para2 = doc.add_paragraph()
    query_para2.style = 'Quote'
    run = query_para2.add_run(query2)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 200)
    
    # ========================================
    # DATA FLOW
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '6. Data Flow: From Source to Database', 1)
    
    doc.add_paragraph('ETL Pipeline Process:', style='Heading 3')
    
    flow_steps = [
        ('STEP 1: EXTRACT', 'Download data from GENESIS API (Regional Database Germany)'),
        ('', 'Source Table: e.g., 13111-07-05-4 (Employment by sector)'),
        ('', 'Output: Raw CSV data (~7,500 rows per year)'),
        ('', ''),
        ('STEP 2: TRANSFORM', 'Parse CSV and clean data'),
        ('', 'Filter to NRW regions only'),
        ('', 'Map region_code → geo_id (lookup in dim_geography)'),
        ('', 'Map year → time_id (lookup in dim_time)'),
        ('', 'Assign indicator_id'),
        ('', 'Validate and clean values'),
        ('', ''),
        ('STEP 3: LOAD', 'Insert into appropriate fact table'),
        ('', 'Bulk insert for performance'),
        ('', 'Validate foreign key constraints'),
        ('', 'Update table registry'),
    ]
    
    for step, desc in flow_steps:
        if step:
            p = doc.add_paragraph(step)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        else:
            doc.add_paragraph(desc, style='List Bullet' if desc else 'Normal')
    
    # ========================================
    # CURRENT STATUS
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '7. Current Database Status', 1)
    
    doc.add_paragraph('As of December 17, 2025:', style='Heading 3')
    
    status_headers = ['Component', 'Status', 'Details']
    status_rows = [
        ('Dimension Tables', 'Complete', '4 tables fully functional'),
        ('dim_geography', 'Populated', '60 NRW regions loaded'),
        ('dim_time', 'Populated', '17 years (2008-2024)'),
        ('dim_indicator', 'Partial', '4 of 40+ indicators defined'),
        ('', '', ''),
        ('Fact Tables', 'In Progress', ''),
        ('fact_demographics', 'Active', '45,000+ records'),
        ('fact_labor_market', 'Pending', 'Using demographics table'),
        ('Other fact tables', 'Pending', 'Schema created, not populated'),
    ]
    
    create_table(doc, status_headers, status_rows)
    
    doc.add_paragraph()
    
    progress_headers = ['Category', 'Tables', 'Status']
    progress_rows = [
        ('Demographics', '1 table', '100% complete (17,556 records)'),
        ('Labor Market', '3 tables', '3 completed, 9 pending'),
        ('Economic Activity', '8 tables', 'Not started'),
        ('Healthcare', '6 tables', 'Not started'),
        ('Public Finance', '3 tables', 'Not started'),
        ('Infrastructure', '1 table', 'Not started'),
        ('Mobility/Commuters', '2 tables', 'Not started'),
    ]
    
    create_table(doc, progress_headers, progress_rows)
    
    # ========================================
    # FAQ
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '8. Frequently Asked Questions', 1)
    
    faqs = [
        ('Why use a star schema instead of separate tables?',
         'Star schema provides flexibility and consistency. Adding new indicators doesn\'t '
         'require creating new tables, just new rows in dim_indicator. All queries follow '
         'the same pattern, making the database easier to learn and use.'),
        
        ('How do I know which fact table to query?',
         'Check the indicator_category in dim_indicator. Demographics → fact_demographics, '
         'Labor Market → fact_labor_market (or fact_demographics currently), '
         'Business → fact_business_economy.'),
        
        ('What\'s the difference between geo_id and region_code?',
         'geo_id is the internal database key (1, 2, 3...), while region_code is the '
         'official GENESIS code (05112, 05, DG). Use region_code in WHERE clauses for '
         'readability.'),
        
        ('Can I add calculated/derived indicators?',
         'Yes! Insert into dim_indicator with is_derived = TRUE, then calculate and insert '
         'the values into the appropriate fact table.'),
        
        ('How is data quality tracked?',
         'Each fact record has a data_quality_flag (V=Validated, E=Estimated, P=Provisional) '
         'and optional confidence_score. Check these fields when data accuracy is critical.'),
    ]
    
    for question, answer in faqs:
        doc.add_paragraph(question, style='Heading 3')
        doc.add_paragraph(answer)
        doc.add_paragraph()
    
    # ========================================
    # GETTING STARTED
    # ========================================
    add_page_break(doc)
    add_styled_heading(doc, '9. Getting Started', 1)
    
    add_styled_heading(doc, 'Step 1: Connect to Database', 2)
    
    conn_text = """
Database: regional_economics
Host: localhost (or your server address)
Port: 5432
Username: (your username)
Password: (your password)
    """
    
    p = doc.add_paragraph(conn_text)
    p.style = 'Quote'
    
    add_styled_heading(doc, 'Step 2: Explore the Data', 2)
    
    explore_queries = [
        ('Count total records', 'SELECT COUNT(*) FROM fact_demographics;'),
        ('List available regions', 'SELECT region_code, region_name FROM dim_geography ORDER BY region_name;'),
        ('List loaded indicators', 'SELECT indicator_id, indicator_code, indicator_name FROM dim_indicator;'),
        ('Check year coverage', 'SELECT DISTINCT year FROM dim_time ORDER BY year;'),
    ]
    
    for desc, query in explore_queries:
        doc.add_paragraph(desc + ':', style='Heading 3')
        p = doc.add_paragraph(query)
        p.runs[0].font.name = 'Consolas'
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0, 0, 200)
    
    # ========================================
    # FOOTER
    # ========================================
    add_page_break(doc)
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = footer_para.add_run('Regional Economics Database for NRW')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Database Structure Guide v1.0')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('For questions or support, contact: Kanyuchi')
    run.font.size = Pt(10)
    run.font.italic = True
    
    # Save document
    output_path = Path('DATABASE_STRUCTURE_GUIDE.docx')
    doc.save(output_path)
    print(f'Word document created successfully: {output_path.absolute()}')
    print(f'File size: {output_path.stat().st_size / 1024:.1f} KB')
    

if __name__ == '__main__':
    main()
